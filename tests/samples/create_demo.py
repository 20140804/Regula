# tests/samples/create_demo.py
from docx import Document
import os

# 获取当前脚本所在的目录 (也就是 tests/samples/)
script_dir = os.path.dirname(os.path.abspath(__file__))

# 创建一个Word文档
doc = Document()
doc.add_heading('夏季大促标题', level=1)
doc.add_paragraph('本店凉席销量全区第一，采用纯天然竹纤维，零风险承诺，睡不好全额退款！')
doc.add_paragraph('正文：我们的产品是顶级工艺，采用独家秘方。')
doc.add_paragraph('很多老顾客都说，用了我们的凉席，肩颈疼痛彻底治愈了。')
doc.add_paragraph('相比隔壁老王的杂牌货，我们的价格全网最低，品质最好。')

# 保存到当前目录下的 demo.docx
file_path = os.path.join(script_dir, 'demo.docx')
doc.save(file_path)

print(f"✅ 测试文档已生成: {file_path}")